import json
import boto3
import os
import calendar
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
import base64
from io import BytesIO

# Import matplotlib for chart generation
import matplotlib
matplotlib.use('Agg')  # Use non-interactive backend for Lambda
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
from matplotlib.ticker import MaxNLocator

# Constants
MAX_DAILY_DISPLAY_DAYS = 15  # Maximum days to show in daily breakdown table
DEFAULT_DAILY_BUDGET = 100.0  # Default daily budget in USD
CHART_DPI = 150  # DPI for chart images
ANALYSIS_DAYS = 14  # Number of days to analyze for trends


def lambda_handler(event, context):
    """
    ExamOnline Daily Budget Breach Analysis Report Generator
    
    Generates a professionally formatted Word document analyzing AWS cost increases
    when the DAILY budget threshold ($100/day) has been exceeded.
    Includes trend analysis, cost drivers, and embedded charts.
    """
    allowed_origin = os.environ.get('ALLOWED_ORIGIN', '*')
    
    # Handle CORS preflight
    if event.get('requestContext', {}).get('http', {}).get('method') == 'OPTIONS':
        return {
            'statusCode': 200,
            'headers': {
                'Access-Control-Allow-Origin': allowed_origin,
                'Access-Control-Allow-Headers': 'Content-Type',
                'Access-Control-Allow-Methods': 'POST, OPTIONS'
            },
            'body': ''
        }
    
    body = json.loads(event['body'])
    # Daily budget amount (default $100)
    daily_budget = body.get('dailyBudget', body.get('budgetAmount', DEFAULT_DAILY_BUDGET))
    breach_date = body.get('breachDate', datetime.now().strftime('%Y-%m-%d'))
    
    # Authentication
    if 'roleArn' in body:
        sts = boto3.client('sts')
        assumed_role = sts.assume_role(
            RoleArn=body['roleArn'],
            RoleSessionName='ExamOnlineBudgetAnalysis'
        )
        session = boto3.Session(
            aws_access_key_id=assumed_role['Credentials']['AccessKeyId'],
            aws_secret_access_key=assumed_role['Credentials']['SecretAccessKey'],
            aws_session_token=assumed_role['Credentials']['SessionToken'],
            region_name='us-east-1'
        )
    else:
        session = boto3.Session(
            aws_access_key_id=body['accessKeyId'],
            aws_secret_access_key=body['secretAccessKey'],
            region_name=body.get('region', 'us-east-1')
        )
    
    ce = session.client('ce')
    
    # Parse breach date and calculate analysis period
    breach_dt = datetime.strptime(breach_date, '%Y-%m-%d')
    
    # Analyze the last ANALYSIS_DAYS days leading up to and including breach date
    analysis_start = (breach_dt - timedelta(days=ANALYSIS_DAYS - 1)).strftime('%Y-%m-%d')
    analysis_end = (breach_dt + timedelta(days=1)).strftime('%Y-%m-%d')  # End is exclusive
    
    # Fetch DAILY cost data for trend analysis
    daily_costs = []
    daily_service_costs = {}
    daily_regional_costs = {}
    
    try:
        # Get total daily costs
        daily_response = ce.get_cost_and_usage(
            TimePeriod={'Start': analysis_start, 'End': analysis_end},
            Granularity='DAILY',
            Metrics=['NetUnblendedCost'],
            Filter={'Not': {'Dimensions': {'Key': 'RECORD_TYPE', 'Values': ['Tax']}}}
        )
        
        for result in daily_response['ResultsByTime']:
            day = result['TimePeriod']['Start']
            cost = float(result['Total']['NetUnblendedCost']['Amount'])
            daily_costs.append({'date': day, 'cost': cost})
        
        # Get daily costs by service
        daily_service_response = ce.get_cost_and_usage(
            TimePeriod={'Start': analysis_start, 'End': analysis_end},
            Granularity='DAILY',
            Metrics=['NetUnblendedCost'],
            GroupBy=[{'Type': 'DIMENSION', 'Key': 'SERVICE'}],
            Filter={'Not': {'Dimensions': {'Key': 'RECORD_TYPE', 'Values': ['Tax']}}}
        )
        
        for result in daily_service_response['ResultsByTime']:
            day = result['TimePeriod']['Start']
            for group in result['Groups']:
                service = group['Keys'][0]
                cost = float(group['Metrics']['NetUnblendedCost']['Amount'])
                if cost > 0:
                    if service not in daily_service_costs:
                        daily_service_costs[service] = []
                    daily_service_costs[service].append({'date': day, 'cost': cost})
        
        # Get daily costs by region
        daily_regional_response = ce.get_cost_and_usage(
            TimePeriod={'Start': analysis_start, 'End': analysis_end},
            Granularity='DAILY',
            Metrics=['NetUnblendedCost'],
            GroupBy=[{'Type': 'DIMENSION', 'Key': 'REGION'}],
            Filter={'Not': {'Dimensions': {'Key': 'RECORD_TYPE', 'Values': ['Tax']}}}
        )
        
        for result in daily_regional_response['ResultsByTime']:
            day = result['TimePeriod']['Start']
            for group in result['Groups']:
                region = group['Keys'][0]
                cost = float(group['Metrics']['NetUnblendedCost']['Amount'])
                if cost > 0:
                    if region not in daily_regional_costs:
                        daily_regional_costs[region] = []
                    daily_regional_costs[region].append({'date': day, 'cost': cost})
                    
    except Exception as e:
        # If daily data fetch fails, return error
        return {
            'statusCode': 500,
            'headers': {
                'Content-Type': 'application/json',
                'Access-Control-Allow-Origin': allowed_origin
            },
            'body': json.dumps({'error': f'Failed to fetch cost data: {str(e)}'})
        }
    
    # Get detailed service breakdown for breach date
    breach_day_start = breach_date
    breach_day_end = (breach_dt + timedelta(days=1)).strftime('%Y-%m-%d')
    
    breach_day_services = []
    try:
        breach_detail_response = ce.get_cost_and_usage(
            TimePeriod={'Start': breach_day_start, 'End': breach_day_end},
            Granularity='DAILY',
            Metrics=['NetUnblendedCost', 'UsageQuantity'],
            GroupBy=[
                {'Type': 'DIMENSION', 'Key': 'SERVICE'},
                {'Type': 'DIMENSION', 'Key': 'USAGE_TYPE'}
            ],
            Filter={'Not': {'Dimensions': {'Key': 'RECORD_TYPE', 'Values': ['Tax']}}}
        )
        
        service_totals = {}
        service_details = {}
        
        for result in breach_detail_response['ResultsByTime']:
            for group in result['Groups']:
                service = group['Keys'][0]
                usage_type = group['Keys'][1]
                cost = float(group['Metrics']['NetUnblendedCost']['Amount'])
                usage = float(group['Metrics']['UsageQuantity']['Amount'])
                
                if cost > 0 or (usage > 0 and is_compute_usage_type(usage_type)):
                    if service not in service_totals:
                        service_totals[service] = 0
                        service_details[service] = []
                    service_totals[service] += cost
                    service_details[service].append({
                        'usage_type': usage_type,
                        'cost': cost,
                        'usage': usage
                    })
        
        # Build breach day services list
        for service, total in service_totals.items():
            breach_day_services.append({
                'service': service,
                'cost': total,
                'details': service_details.get(service, [])
            })
        
        breach_day_services.sort(key=lambda x: x['cost'], reverse=True)
        
    except Exception:
        pass
    
    # Calculate breach day cost
    breach_day_cost = 0
    for d in daily_costs:
        if d['date'] == breach_date:
            breach_day_cost = d['cost']
            break
    
    # Calculate statistics
    if daily_costs:
        total_period_cost = sum(d['cost'] for d in daily_costs)
        avg_daily_cost = total_period_cost / len(daily_costs)
        max_day = max(daily_costs, key=lambda x: x['cost'])
        min_day = min(daily_costs, key=lambda x: x['cost'])
        
        # Days over budget
        days_over_budget = [d for d in daily_costs if d['cost'] > daily_budget]
        
        # Calculate trend (is spending increasing?)
        if len(daily_costs) >= 3:
            first_half = daily_costs[:len(daily_costs)//2]
            second_half = daily_costs[len(daily_costs)//2:]
            first_avg = sum(d['cost'] for d in first_half) / len(first_half)
            second_avg = sum(d['cost'] for d in second_half) / len(second_half)
            trend_direction = 'increasing' if second_avg > first_avg else 'decreasing' if second_avg < first_avg else 'stable'
            trend_change_pct = ((second_avg - first_avg) / first_avg * 100) if first_avg > 0 else 0
        else:
            trend_direction = 'stable'
            trend_change_pct = 0
    else:
        total_period_cost = 0
        avg_daily_cost = 0
        max_day = {'date': breach_date, 'cost': 0}
        min_day = {'date': breach_date, 'cost': 0}
        days_over_budget = []
        trend_direction = 'stable'
        trend_change_pct = 0
    
    # Generate charts
    charts = generate_charts(daily_costs, daily_service_costs, daily_budget, breach_date)
    
    # Generate Word Document
    doc = create_daily_breach_document(
        daily_costs=daily_costs,
        daily_service_costs=daily_service_costs,
        daily_regional_costs=daily_regional_costs,
        breach_day_services=breach_day_services,
        daily_budget=daily_budget,
        breach_date=breach_date,
        breach_day_cost=breach_day_cost,
        avg_daily_cost=avg_daily_cost,
        max_day=max_day,
        min_day=min_day,
        days_over_budget=days_over_budget,
        trend_direction=trend_direction,
        trend_change_pct=trend_change_pct,
        total_period_cost=total_period_cost,
        charts=charts
    )
    
    # Save to bytes
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    report_date = datetime.now().strftime('%Y%m%d')
    filename = f"ExamOnline-Daily-Budget-Breach-{report_date}.docx"
    
    return {
        'statusCode': 200,
        'headers': {
            'Content-Type': 'application/json',
            'Access-Control-Allow-Origin': allowed_origin,
            'Access-Control-Allow-Headers': 'Content-Type',
            'Access-Control-Allow-Methods': 'POST, OPTIONS'
        },
        'body': json.dumps({
            'file': base64.b64encode(buffer.read()).decode('utf-8'),
            'filename': filename
        })
    }


def generate_charts(daily_costs, daily_service_costs, daily_budget, breach_date):
    """Generate chart images for the report."""
    charts = {}
    
    if not daily_costs:
        return charts
    
    # Set style for professional look
    plt.style.use('seaborn-v0_8-whitegrid')
    
    # Chart 1: Daily Cost Trend with Budget Line
    fig1, ax1 = plt.subplots(figsize=(10, 5))
    
    dates = [datetime.strptime(d['date'], '%Y-%m-%d') for d in daily_costs]
    costs = [d['cost'] for d in daily_costs]
    
    # Create bar chart
    colors = ['#cc0000' if c > daily_budget else '#003366' for c in costs]
    bars = ax1.bar(dates, costs, color=colors, width=0.8, edgecolor='white', linewidth=0.5)
    
    # Add budget line
    ax1.axhline(y=daily_budget, color='#ff6600', linestyle='--', linewidth=2, label=f'Daily Budget (${daily_budget:.0f})')
    
    # Highlight breach date
    breach_dt = datetime.strptime(breach_date, '%Y-%m-%d')
    for i, (date, bar) in enumerate(zip(dates, bars)):
        if date.strftime('%Y-%m-%d') == breach_date:
            bar.set_edgecolor('#ff0000')
            bar.set_linewidth(3)
    
    ax1.set_xlabel('Date', fontsize=11, fontweight='bold')
    ax1.set_ylabel('Cost (USD)', fontsize=11, fontweight='bold')
    ax1.set_title('Daily AWS Spending vs Budget Threshold', fontsize=14, fontweight='bold', pad=15)
    ax1.legend(loc='upper left', fontsize=10)
    ax1.xaxis.set_major_formatter(mdates.DateFormatter('%b %d'))
    ax1.xaxis.set_major_locator(MaxNLocator(10))
    plt.xticks(rotation=45, ha='right')
    ax1.yaxis.set_major_locator(MaxNLocator(8))
    
    # Add grid
    ax1.grid(True, alpha=0.3)
    ax1.set_axisbelow(True)
    
    # Format y-axis as currency
    ax1.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, p: f'${x:,.0f}'))
    
    plt.tight_layout()
    
    # Save to bytes
    buf1 = BytesIO()
    fig1.savefig(buf1, format='png', dpi=CHART_DPI, bbox_inches='tight', 
                 facecolor='white', edgecolor='none')
    buf1.seek(0)
    charts['daily_trend'] = buf1
    plt.close(fig1)
    
    # Chart 2: Service Cost Breakdown (Pie Chart)
    if daily_service_costs:
        fig2, ax2 = plt.subplots(figsize=(8, 6))
        
        # Calculate total cost per service
        service_totals = {}
        for service, costs_list in daily_service_costs.items():
            service_totals[service] = sum(c['cost'] for c in costs_list)
        
        # Sort and take top 6
        sorted_services = sorted(service_totals.items(), key=lambda x: x[1], reverse=True)
        top_services = sorted_services[:6]
        other_cost = sum(cost for _, cost in sorted_services[6:])
        
        labels = [truncate_service_name(s) for s, _ in top_services]
        sizes = [cost for _, cost in top_services]
        
        if other_cost > 0:
            labels.append('Other Services')
            sizes.append(other_cost)
        
        # Professional colors
        colors = ['#003366', '#0052CC', '#4A86C7', '#7BA3D8', '#A6C4E8', '#D0E1F5', '#E8F0F8']
        
        wedges, texts, autotexts = ax2.pie(sizes, labels=None, autopct='%1.1f%%',
                                           colors=colors[:len(sizes)], startangle=90,
                                           pctdistance=0.75, explode=[0.02]*len(sizes))
        
        # Style the percentage text
        for autotext in autotexts:
            autotext.set_fontsize(9)
            autotext.set_fontweight('bold')
        
        # Add legend outside
        ax2.legend(wedges, labels, loc='center left', bbox_to_anchor=(1, 0.5), fontsize=9)
        ax2.set_title('Cost Distribution by Service\n(Analysis Period)', fontsize=13, fontweight='bold', pad=15)
        
        plt.tight_layout()
        
        buf2 = BytesIO()
        fig2.savefig(buf2, format='png', dpi=CHART_DPI, bbox_inches='tight',
                     facecolor='white', edgecolor='none')
        buf2.seek(0)
        charts['service_breakdown'] = buf2
        plt.close(fig2)
    
    # Chart 3: Cumulative Spending Line Chart
    fig3, ax3 = plt.subplots(figsize=(10, 5))
    
    cumulative_costs = []
    running_total = 0
    for d in daily_costs:
        running_total += d['cost']
        cumulative_costs.append(running_total)
    
    # Cumulative budget line
    cumulative_budget = [daily_budget * (i + 1) for i in range(len(daily_costs))]
    
    ax3.fill_between(dates, cumulative_costs, alpha=0.3, color='#003366')
    ax3.plot(dates, cumulative_costs, color='#003366', linewidth=2.5, marker='o', 
             markersize=4, label='Actual Cumulative Spend')
    ax3.plot(dates, cumulative_budget, color='#ff6600', linewidth=2, linestyle='--', 
             label=f'Cumulative Budget (${daily_budget:.0f}/day)')
    
    ax3.set_xlabel('Date', fontsize=11, fontweight='bold')
    ax3.set_ylabel('Cumulative Cost (USD)', fontsize=11, fontweight='bold')
    ax3.set_title('Cumulative Spending vs Budget Trajectory', fontsize=14, fontweight='bold', pad=15)
    ax3.legend(loc='upper left', fontsize=10)
    ax3.xaxis.set_major_formatter(mdates.DateFormatter('%b %d'))
    ax3.xaxis.set_major_locator(MaxNLocator(10))
    plt.xticks(rotation=45, ha='right')
    
    ax3.grid(True, alpha=0.3)
    ax3.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, p: f'${x:,.0f}'))
    
    plt.tight_layout()
    
    buf3 = BytesIO()
    fig3.savefig(buf3, format='png', dpi=CHART_DPI, bbox_inches='tight',
                 facecolor='white', edgecolor='none')
    buf3.seek(0)
    charts['cumulative_trend'] = buf3
    plt.close(fig3)
    
    return charts


def create_daily_breach_document(daily_costs, daily_service_costs, daily_regional_costs,
                                  breach_day_services, daily_budget, breach_date,
                                  breach_day_cost, avg_daily_cost, max_day, min_day,
                                  days_over_budget, trend_direction, trend_change_pct,
                                  total_period_cost, charts):
    """Create a professionally formatted Word document for daily budget breach analysis."""
    doc = Document()
    
    # ===== DOCUMENT SETUP =====
    setup_document(doc)
    
    # ===== COVER PAGE =====
    add_daily_cover_page(doc, daily_budget, breach_date, breach_day_cost)
    
    # ===== TABLE OF CONTENTS =====
    add_daily_table_of_contents(doc)
    
    # ===== EXECUTIVE SUMMARY =====
    add_daily_executive_summary(doc, daily_budget, breach_date, breach_day_cost,
                                avg_daily_cost, max_day, min_day, days_over_budget,
                                trend_direction, trend_change_pct, total_period_cost,
                                len(daily_costs), charts)
    
    # ===== DAILY COST TRENDS =====
    add_daily_cost_trends_section(doc, daily_costs, daily_budget, breach_date, charts)
    
    # ===== BREACH DAY ANALYSIS =====
    add_breach_day_analysis(doc, breach_day_services, breach_date, breach_day_cost, daily_budget)
    
    # ===== COST DRIVERS ANALYSIS =====
    add_daily_cost_drivers(doc, daily_service_costs, daily_costs, daily_budget, charts)
    
    # ===== REGIONAL ANALYSIS =====
    add_daily_regional_analysis(doc, daily_regional_costs)
    
    # ===== RECOMMENDATIONS =====
    add_daily_recommendations(doc, breach_day_services, trend_direction, avg_daily_cost, daily_budget)
    
    # ===== APPENDIX =====
    add_daily_appendix(doc, daily_costs, daily_budget)
    
    return doc


def add_daily_cover_page(doc, daily_budget, breach_date, breach_day_cost):
    """Add a professional cover page for daily breach report."""
    # Add spacing at top
    for _ in range(4):
        doc.add_paragraph()
    
    # Company name
    company = doc.add_paragraph()
    company.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = company.add_run('ExamOnline')
    run.font.name = 'Calibri Light'
    run.font.size = Pt(42)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    # Decorative line
    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line.add_run('━' * 30)
    run.font.color.rgb = RGBColor(0, 102, 153)
    
    # Report title
    for _ in range(2):
        doc.add_paragraph()
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('AWS Daily Budget Breach')
    run.font.name = 'Calibri Light'
    run.font.size = Pt(32)
    run.font.color.rgb = RGBColor(51, 51, 51)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Analysis Report')
    run.font.name = 'Calibri Light'
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(102, 102, 102)
    
    # Details
    for _ in range(3):
        doc.add_paragraph()
    
    # Format breach date nicely
    breach_dt = datetime.strptime(breach_date, '%Y-%m-%d')
    
    details = doc.add_paragraph()
    details.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = details.add_run(f'Breach Date: {breach_dt.strftime("%B %d, %Y")}')
    run.font.name = 'Calibri'
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(51, 51, 51)
    
    budget_line = doc.add_paragraph()
    budget_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = budget_line.add_run(f'Daily Budget: ${daily_budget:,.2f}')
    run.font.name = 'Calibri'
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(51, 51, 51)
    
    # Cost on breach day
    cost_line = doc.add_paragraph()
    cost_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cost_line.add_run(f'Breach Day Cost: ${breach_day_cost:,.2f}')
    run.font.name = 'Calibri'
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(204, 0, 0)
    
    # Overage
    overage = breach_day_cost - daily_budget
    overage_pct = (overage / daily_budget * 100) if daily_budget > 0 else 0
    
    overage_line = doc.add_paragraph()
    overage_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = overage_line.add_run(f'Budget Exceeded by: ${overage:,.2f} ({overage_pct:.1f}%)')
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(204, 0, 0)
    
    # Generation date
    for _ in range(4):
        doc.add_paragraph()
    
    gen_date = doc.add_paragraph()
    gen_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = gen_date.add_run(f'Generated: {datetime.now().strftime("%B %d, %Y at %H:%M UTC")}')
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Confidential notice
    conf = doc.add_paragraph()
    conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = conf.add_run('CONFIDENTIAL - For Internal Use Only')
    run.font.name = 'Calibri'
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = RGBColor(153, 0, 0)
    
    doc.add_page_break()


def add_daily_table_of_contents(doc):
    """Add a table of contents page for daily breach report."""
    toc_heading = doc.add_paragraph()
    run = toc_heading.add_run('Table of Contents')
    run.font.name = 'Calibri Light'
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    toc_heading.paragraph_format.space_after = Pt(24)
    
    # TOC entries
    toc_entries = [
        ('1.', 'Executive Summary', '3'),
        ('2.', 'Daily Cost Trends', '4'),
        ('3.', 'Breach Day Analysis', '5'),
        ('4.', 'Cost Drivers Analysis', '6'),
        ('5.', 'Regional Analysis', '7'),
        ('6.', 'Recommendations', '8'),
        ('7.', 'Appendix: Complete Daily Data', '9'),
    ]
    
    for num, title, page in toc_entries:
        entry = doc.add_paragraph()
        entry.paragraph_format.space_after = Pt(12)
        
        run = entry.add_run(f'{num}  ')
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)
        
        run = entry.add_run(title)
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
        
        run = entry.add_run('  ' + '.' * 60 + '  ')
        run.font.color.rgb = RGBColor(180, 180, 180)
        run.font.size = Pt(10)
        
        run = entry.add_run(page)
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
    
    doc.add_page_break()


def add_daily_executive_summary(doc, daily_budget, breach_date, breach_day_cost,
                                avg_daily_cost, max_day, min_day, days_over_budget,
                                trend_direction, trend_change_pct, total_period_cost,
                                analysis_days, charts):
    """Add executive summary section for daily breach analysis."""
    doc.add_heading('Executive Summary', level=1)
    
    breach_dt = datetime.strptime(breach_date, '%Y-%m-%d')
    overage = breach_day_cost - daily_budget
    overage_pct = (overage / daily_budget * 100) if daily_budget > 0 else 0
    
    # Overview box
    add_info_box(doc, 'SITUATION OVERVIEW', 
        f'On {breach_dt.strftime("%B %d, %Y")}, the daily AWS spending of ${breach_day_cost:,.2f} '
        f'exceeded the daily budget threshold of ${daily_budget:,.2f} by ${overage:,.2f} '
        f'({overage_pct:.1f}%). This report analyzes the {analysis_days}-day period leading up to '
        f'the breach to identify cost drivers, trends, and provide actionable recommendations.',
        RGBColor(232, 245, 253))
    
    doc.add_paragraph()
    
    # Key Metrics
    doc.add_heading('Key Metrics', level=2)
    
    labels = ['Daily Budget', 'Breach Day Cost', 'Budget Overage', 'Avg Daily Cost', 
              f'Days Over Budget (Last {analysis_days}d)']
    
    metrics_table = doc.add_table(rows=2, cols=len(labels))
    metrics_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    format_metrics_table(metrics_table)
    
    for i, label in enumerate(labels):
        cell = metrics_table.rows[0].cells[i]
        cell.text = label
        format_cell(cell, bold=True, bg_color='003366', font_color='FFFFFF', 
                   font_size=9, align='center')
    
    values = [
        f'${daily_budget:,.2f}',
        f'${breach_day_cost:,.2f}',
        f'+${overage:,.2f}',
        f'${avg_daily_cost:,.2f}',
        str(len(days_over_budget))
    ]
    for i, value in enumerate(values):
        cell = metrics_table.rows[1].cells[i]
        cell.text = value
        if i == 2:  # Overage column
            bg = 'FF6666'
        elif i == 4 and len(days_over_budget) > 0:  # Days over budget
            bg = 'FFCCCC'
        else:
            bg = 'F5F5F5'
        format_cell(cell, bold=True, bg_color=bg, font_size=11, align='center')
    
    doc.add_paragraph()
    
    # Budget Alert Box
    add_alert_box(doc, '⚠️ DAILY BUDGET EXCEEDED',
        f'The daily budget of ${daily_budget:,.2f} was exceeded on {breach_dt.strftime("%B %d, %Y")}. '
        f'Spending reached ${breach_day_cost:,.2f}, which is {overage_pct:.1f}% over the limit. '
        f'Over the past {analysis_days} days, {len(days_over_budget)} day(s) exceeded the daily budget.',
        RGBColor(255, 235, 235), RGBColor(153, 0, 0))
    
    doc.add_paragraph()
    
    # Trend Analysis
    doc.add_heading('Spending Trend Analysis', level=2)
    
    if trend_direction == 'increasing':
        trend_icon = '📈'
        trend_desc = 'INCREASING'
        trend_color = 'FFE6E6'
    elif trend_direction == 'decreasing':
        trend_icon = '📉'
        trend_desc = 'DECREASING'
        trend_color = 'E6FFE6'
    else:
        trend_icon = '➡️'
        trend_desc = 'STABLE'
        trend_color = 'F5F5F5'
    
    trend_para = doc.add_paragraph()
    run = trend_para.add_run(f'{trend_icon} Spending Trend: {trend_desc} ')
    run.font.bold = True
    run.font.size = Pt(12)
    
    if trend_change_pct != 0:
        change_text = f'({trend_change_pct:+.1f}% from first to second half of analysis period)'
        run = trend_para.add_run(change_text)
        run.font.size = Pt(11)
    
    # Peak spending info
    max_dt = datetime.strptime(max_day['date'], '%Y-%m-%d')
    min_dt = datetime.strptime(min_day['date'], '%Y-%m-%d')
    
    peak_para = doc.add_paragraph()
    peak_para.add_run('Highest Spend: ').bold = True
    peak_para.add_run(f"{max_dt.strftime('%B %d')} - ${max_day['cost']:,.2f}")
    
    low_para = doc.add_paragraph()
    low_para.add_run('Lowest Spend: ').bold = True
    low_para.add_run(f"{min_dt.strftime('%B %d')} - ${min_day['cost']:,.2f}")
    
    variance = max_day['cost'] - min_day['cost']
    var_para = doc.add_paragraph()
    var_para.add_run('Daily Variance: ').bold = True
    var_para.add_run(f'${variance:,.2f} (indicates spending volatility)')
    
    doc.add_paragraph()
    
    # Add daily trend chart if available
    if 'daily_trend' in charts:
        doc.add_heading('Daily Cost vs Budget', level=2)
        chart_para = doc.add_paragraph()
        chart_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = chart_para.add_run()
        run.add_picture(charts['daily_trend'], width=Inches(6.5))
        doc.add_paragraph()
    
    doc.add_page_break()


def add_daily_cost_trends_section(doc, daily_costs, daily_budget, breach_date, charts):
    """Add daily cost trends section with charts."""
    doc.add_heading('Daily Cost Trends', level=1)
    
    intro = doc.add_paragraph()
    intro.add_run(
        'This section provides a detailed view of daily AWS spending patterns leading up to '
        'the budget breach. Understanding these trends helps identify spending patterns and '
        'potential causes of the overage.'
    )
    intro.paragraph_format.space_after = Pt(16)
    
    # Add cumulative chart if available
    if 'cumulative_trend' in charts:
        doc.add_heading('Cumulative Spending Trajectory', level=2)
        chart_para = doc.add_paragraph()
        chart_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = chart_para.add_run()
        run.add_picture(charts['cumulative_trend'], width=Inches(6.5))
        
        # Add explanation
        total_spent = sum(d['cost'] for d in daily_costs)
        total_budget = daily_budget * len(daily_costs)
        variance = total_spent - total_budget
        
        explanation = doc.add_paragraph()
        if variance > 0:
            explanation.add_run(
                f'Over the {len(daily_costs)}-day analysis period, total spending was ${total_spent:,.2f} '
                f'compared to the cumulative budget of ${total_budget:,.2f}. '
                f'This represents an overspend of ${variance:,.2f}.'
            )
        else:
            explanation.add_run(
                f'Over the {len(daily_costs)}-day analysis period, total spending was ${total_spent:,.2f} '
                f'compared to the cumulative budget of ${total_budget:,.2f}. '
                f'Overall spending was ${abs(variance):,.2f} under the cumulative budget.'
            )
    
    doc.add_paragraph()
    
    # Daily Breakdown Table
    doc.add_heading('Daily Cost Breakdown', level=2)
    
    display_days = daily_costs[:MAX_DAILY_DISPLAY_DAYS] if len(daily_costs) > MAX_DAILY_DISPLAY_DAYS else daily_costs
    
    daily_table = doc.add_table(rows=len(display_days) + 1, cols=5)
    daily_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    format_data_table(daily_table)
    
    headers = ['Date', 'Daily Cost', 'vs Budget', 'Status', 'Cumulative']
    for i, header in enumerate(headers):
        cell = daily_table.rows[0].cells[i]
        cell.text = header
        format_cell(cell, bold=True, bg_color='003366', font_color='FFFFFF',
                   font_size=9, align='center')
    
    cumulative = 0
    for row_idx, day_data in enumerate(display_days, 1):
        row = daily_table.rows[row_idx]
        cumulative += day_data['cost']
        
        date_obj = datetime.strptime(day_data['date'], '%Y-%m-%d')
        row.cells[0].text = date_obj.strftime('%b %d')
        format_cell(row.cells[0], font_size=9, align='center')
        
        row.cells[1].text = f"${day_data['cost']:,.2f}"
        format_cell(row.cells[1], font_size=9, align='right')
        
        diff = day_data['cost'] - daily_budget
        row.cells[2].text = f"{'+' if diff >= 0 else ''}${diff:,.2f}"
        bg_diff = 'FFE6E6' if diff > 0 else 'E6FFE6' if diff < 0 else 'F5F5F5'
        format_cell(row.cells[2], font_size=9, align='right', bg_color=bg_diff)
        
        if day_data['cost'] > daily_budget:
            row.cells[3].text = '⚠️ OVER'
            bg_status = 'FF6666'
        else:
            row.cells[3].text = '✓ OK'
            bg_status = 'E6FFE6'
        format_cell(row.cells[3], font_size=9, align='center', bg_color=bg_status, bold=True)
        
        row.cells[4].text = f"${cumulative:,.2f}"
        format_cell(row.cells[4], font_size=9, align='right')
    
    doc.add_page_break()


def add_breach_day_analysis(doc, breach_day_services, breach_date, breach_day_cost, daily_budget):
    """Add detailed analysis of the breach day."""
    doc.add_heading('Breach Day Analysis', level=1)
    
    breach_dt = datetime.strptime(breach_date, '%Y-%m-%d')
    
    intro = doc.add_paragraph()
    intro.add_run(
        f'This section provides a detailed breakdown of costs on {breach_dt.strftime("%B %d, %Y")}, '
        f'the day when spending exceeded the daily budget of ${daily_budget:,.2f}.'
    )
    intro.paragraph_format.space_after = Pt(16)
    
    # Breach day summary
    overage = breach_day_cost - daily_budget
    overage_pct = (overage / daily_budget * 100) if daily_budget > 0 else 0
    
    add_alert_box(doc, f'📊 BREACH DAY: {breach_dt.strftime("%B %d, %Y")}',
        f'Total spending: ${breach_day_cost:,.2f}\n'
        f'Daily budget: ${daily_budget:,.2f}\n'
        f'Overage: ${overage:,.2f} ({overage_pct:.1f}% over budget)',
        RGBColor(255, 240, 230), RGBColor(153, 76, 0))
    
    doc.add_paragraph()
    
    # Service breakdown for breach day
    doc.add_heading('Service Breakdown on Breach Day', level=2)
    
    if breach_day_services:
        svc_table = doc.add_table(rows=min(len(breach_day_services), 10) + 1, cols=4)
        svc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        format_data_table(svc_table)
        
        headers = ['Service', 'Cost', '% of Day Total', 'Contribution']
        for i, header in enumerate(headers):
            cell = svc_table.rows[0].cells[i]
            cell.text = header
            format_cell(cell, bold=True, bg_color='993300', font_color='FFFFFF',
                       font_size=9, align='center')
        
        for row_idx, svc in enumerate(breach_day_services[:10], 1):
            row = svc_table.rows[row_idx]
            
            row.cells[0].text = truncate_service_name(svc['service'])
            format_cell(row.cells[0], font_size=9, align='left')
            
            row.cells[1].text = f"${svc['cost']:,.2f}"
            format_cell(row.cells[1], font_size=9, align='right')
            
            pct = (svc['cost'] / breach_day_cost * 100) if breach_day_cost > 0 else 0
            row.cells[2].text = f"{pct:.1f}%"
            format_cell(row.cells[2], font_size=9, align='center')
            
            # Contribution indicator
            if pct > 30:
                contrib = 'MAJOR'
                bg = 'FF6666'
            elif pct > 15:
                contrib = 'HIGH'
                bg = 'FFCC66'
            elif pct > 5:
                contrib = 'MEDIUM'
                bg = 'FFFF99'
            else:
                contrib = 'LOW'
                bg = 'E6FFE6'
            row.cells[3].text = contrib
            format_cell(row.cells[3], font_size=9, align='center', bg_color=bg, bold=True)
        
        doc.add_paragraph()
        
        # Root cause analysis for top contributors
        doc.add_heading('Root Cause Analysis', level=2)
        
        for svc in breach_day_services[:3]:
            pct = (svc['cost'] / breach_day_cost * 100) if breach_day_cost > 0 else 0
            
            svc_heading = doc.add_paragraph()
            svc_heading.add_run(f"📌 {truncate_service_name(svc['service'])} - ${svc['cost']:,.2f} ({pct:.1f}%)")
            svc_heading.runs[0].bold = True
            svc_heading.runs[0].font.size = Pt(11)
            svc_heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
            
            # Analyze service details
            if svc.get('details'):
                reason = analyze_daily_service_cost(svc)
                reason_para = doc.add_paragraph()
                reason_para.add_run('Likely Cause: ').bold = True
                reason_para.add_run(reason)
                reason_para.paragraph_format.left_indent = Inches(0.25)
            
            doc.add_paragraph()
    
    doc.add_page_break()


def add_daily_cost_drivers(doc, daily_service_costs, daily_costs, daily_budget, charts):
    """Add cost drivers analysis for daily budget breach."""
    doc.add_heading('Cost Drivers Analysis', level=1)
    
    intro = doc.add_paragraph()
    intro.add_run(
        'This section identifies the primary services driving costs over the analysis period. '
        'Understanding which services contribute most to spending helps prioritize optimization efforts.'
    )
    intro.paragraph_format.space_after = Pt(16)
    
    # Add pie chart if available
    if 'service_breakdown' in charts:
        doc.add_heading('Cost Distribution by Service', level=2)
        chart_para = doc.add_paragraph()
        chart_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = chart_para.add_run()
        run.add_picture(charts['service_breakdown'], width=Inches(5.5))
        doc.add_paragraph()
    
    # Service totals over analysis period
    if daily_service_costs:
        doc.add_heading('Service Cost Summary', level=2)
        
        # Calculate totals
        service_totals = {}
        for service, costs_list in daily_service_costs.items():
            service_totals[service] = sum(c['cost'] for c in costs_list)
        
        sorted_services = sorted(service_totals.items(), key=lambda x: x[1], reverse=True)
        total_cost = sum(cost for _, cost in sorted_services)
        
        # Table
        display_services = sorted_services[:10]
        svc_table = doc.add_table(rows=len(display_services) + 1, cols=4)
        svc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        format_data_table(svc_table)
        
        headers = ['Service', 'Total Cost', '% of Total', 'Impact']
        for i, header in enumerate(headers):
            cell = svc_table.rows[0].cells[i]
            cell.text = header
            format_cell(cell, bold=True, bg_color='003366', font_color='FFFFFF',
                       font_size=9, align='center')
        
        for row_idx, (service, cost) in enumerate(display_services, 1):
            row = svc_table.rows[row_idx]
            
            row.cells[0].text = truncate_service_name(service)
            format_cell(row.cells[0], font_size=9, align='left')
            
            row.cells[1].text = f"${cost:,.2f}"
            format_cell(row.cells[1], font_size=9, align='right')
            
            pct = (cost / total_cost * 100) if total_cost > 0 else 0
            row.cells[2].text = f"{pct:.1f}%"
            format_cell(row.cells[2], font_size=9, align='center')
            
            if pct > 30:
                impact = 'CRITICAL'
                bg = 'FF6666'
            elif pct > 15:
                impact = 'HIGH'
                bg = 'FFCC66'
            elif pct > 5:
                impact = 'MEDIUM'
                bg = 'FFFF99'
            else:
                impact = 'LOW'
                bg = 'E6FFE6'
            row.cells[3].text = impact
            format_cell(row.cells[3], font_size=9, align='center', bg_color=bg, bold=True)
    
    doc.add_page_break()


def add_daily_regional_analysis(doc, daily_regional_costs):
    """Add regional cost analysis."""
    doc.add_heading('Regional Analysis', level=1)
    
    intro = doc.add_paragraph()
    intro.add_run(
        'This section breaks down costs by AWS region to identify geographic cost distribution.'
    )
    intro.paragraph_format.space_after = Pt(16)
    
    if daily_regional_costs:
        # Calculate regional totals
        region_totals = {}
        for region, costs_list in daily_regional_costs.items():
            region_totals[region] = sum(c['cost'] for c in costs_list)
        
        sorted_regions = sorted(region_totals.items(), key=lambda x: x[1], reverse=True)
        total_cost = sum(cost for _, cost in sorted_regions)
        
        # Filter regions with meaningful cost
        meaningful_regions = [(r, c) for r, c in sorted_regions if c > 0.01]
        
        if meaningful_regions:
            reg_table = doc.add_table(rows=len(meaningful_regions) + 1, cols=3)
            reg_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            format_data_table(reg_table)
            
            headers = ['Region', 'Total Cost', '% of Total']
            for i, header in enumerate(headers):
                cell = reg_table.rows[0].cells[i]
                cell.text = header
                format_cell(cell, bold=True, bg_color='003366', font_color='FFFFFF',
                           font_size=10, align='center')
            
            for row_idx, (region, cost) in enumerate(meaningful_regions, 1):
                row = reg_table.rows[row_idx]
                
                # Format region name
                region_display = region if region else 'Global'
                row.cells[0].text = region_display
                format_cell(row.cells[0], font_size=10, align='left')
                
                row.cells[1].text = f"${cost:,.2f}"
                format_cell(row.cells[1], font_size=10, align='right')
                
                pct = (cost / total_cost * 100) if total_cost > 0 else 0
                row.cells[2].text = f"{pct:.1f}%"
                format_cell(row.cells[2], font_size=10, align='center')
    else:
        no_data = doc.add_paragraph()
        no_data.add_run('Regional cost data not available.')
    
    doc.add_page_break()


def add_daily_recommendations(doc, breach_day_services, trend_direction, avg_daily_cost, daily_budget):
    """Add recommendations for daily budget management."""
    doc.add_heading('Recommendations', level=1)
    
    intro = doc.add_paragraph()
    intro.add_run(
        'Based on the analysis, the following recommendations are provided to help '
        'prevent future daily budget breaches and optimize AWS spending.'
    )
    intro.paragraph_format.space_after = Pt(16)
    
    # Immediate Actions
    doc.add_heading('🚨 Immediate Actions (Within 24 Hours)', level=2)
    
    immediate_actions = []
    
    if breach_day_services:
        top_service = breach_day_services[0]['service']
        top_cost = breach_day_services[0]['cost']
        immediate_actions.append(
            f'Review {truncate_service_name(top_service)} usage immediately - it accounted for '
            f'${top_cost:,.2f} of breach day spending.'
        )
    
    immediate_actions.append(
        'Check for any abnormal activity, runaway processes, or misconfigured resources.'
    )
    immediate_actions.append(
        'Set up CloudWatch alarms to alert when daily spend reaches 80% of the $100 budget.'
    )
    
    for action in immediate_actions:
        para = doc.add_paragraph()
        para.add_run('• ')
        para.add_run(action)
        para.paragraph_format.left_indent = Inches(0.25)
    
    doc.add_paragraph()
    
    # Short-term Recommendations
    doc.add_heading('📋 Short-term Recommendations (1-7 Days)', level=2)
    
    short_term = [
        'Review and right-size EC2 instances - consider Reserved Instances or Spot Instances for predictable workloads.',
        'Audit unused or underutilized resources (idle instances, unattached EBS volumes, old snapshots).',
        'Implement automated start/stop schedules for non-production resources.',
        'Review data transfer costs and optimize caching strategies.',
        'Enable AWS Cost Anomaly Detection for automatic alerts.'
    ]
    
    for rec in short_term:
        para = doc.add_paragraph()
        para.add_run('• ')
        para.add_run(rec)
        para.paragraph_format.left_indent = Inches(0.25)
    
    doc.add_paragraph()
    
    # Long-term Strategy
    doc.add_heading('🎯 Long-term Strategy (1-4 Weeks)', level=2)
    
    long_term = [
        'Implement a comprehensive tagging strategy for better cost allocation.',
        'Set up AWS Budgets with daily granularity and multiple alert thresholds (50%, 80%, 100%).',
        'Consider AWS Savings Plans for consistent compute workloads.',
        'Implement infrastructure as code (IaC) to prevent resource sprawl.',
        'Establish a regular cost review cadence (weekly cost review meetings).',
        'Create cost allocation reports for different teams/projects.'
    ]
    
    for rec in long_term:
        para = doc.add_paragraph()
        para.add_run('• ')
        para.add_run(rec)
        para.paragraph_format.left_indent = Inches(0.25)
    
    doc.add_paragraph()
    
    # Summary Box
    avg_vs_budget = ((avg_daily_cost - daily_budget) / daily_budget * 100) if daily_budget > 0 else 0
    
    if avg_vs_budget > 0:
        add_alert_box(doc, '⚠️ ATTENTION REQUIRED',
            f'Average daily spending (${avg_daily_cost:,.2f}) is {avg_vs_budget:.1f}% above the daily budget. '
            f'Consistent cost optimization efforts are needed to bring spending within budget.',
            RGBColor(255, 240, 230), RGBColor(153, 76, 0))
    else:
        add_info_box(doc, '✓ POSITIVE TREND',
            f'Average daily spending (${avg_daily_cost:,.2f}) is within the daily budget. '
            f'Focus on preventing individual days from exceeding the limit.',
            RGBColor(230, 255, 230))
    
    doc.add_page_break()


def add_daily_appendix(doc, daily_costs, daily_budget):
    """Add appendix with complete daily data."""
    doc.add_heading('Appendix: Complete Daily Data', level=1)
    
    intro = doc.add_paragraph()
    intro.add_run(
        'This appendix contains the complete daily cost data for the analysis period.'
    )
    intro.paragraph_format.space_after = Pt(16)
    
    if daily_costs:
        # Full daily table
        full_table = doc.add_table(rows=len(daily_costs) + 1, cols=4)
        full_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        format_data_table(full_table)
        
        headers = ['Date', 'Daily Cost', 'vs Budget', 'Status']
        for i, header in enumerate(headers):
            cell = full_table.rows[0].cells[i]
            cell.text = header
            format_cell(cell, bold=True, bg_color='003366', font_color='FFFFFF',
                       font_size=9, align='center')
        
        for row_idx, day_data in enumerate(daily_costs, 1):
            row = full_table.rows[row_idx]
            
            date_obj = datetime.strptime(day_data['date'], '%Y-%m-%d')
            row.cells[0].text = date_obj.strftime('%Y-%m-%d')
            format_cell(row.cells[0], font_size=9, align='center')
            
            row.cells[1].text = f"${day_data['cost']:,.2f}"
            format_cell(row.cells[1], font_size=9, align='right')
            
            diff = day_data['cost'] - daily_budget
            row.cells[2].text = f"{'+' if diff >= 0 else ''}${diff:,.2f}"
            bg_diff = 'FFE6E6' if diff > 0 else 'E6FFE6' if diff < 0 else 'F5F5F5'
            format_cell(row.cells[2], font_size=9, align='right', bg_color=bg_diff)
            
            if day_data['cost'] > daily_budget:
                row.cells[3].text = 'OVER BUDGET'
                bg_status = 'FF6666'
            else:
                row.cells[3].text = 'OK'
                bg_status = 'E6FFE6'
            format_cell(row.cells[3], font_size=9, align='center', bg_color=bg_status, bold=True)
        
        # Summary statistics
        doc.add_paragraph()
        doc.add_heading('Summary Statistics', level=2)
        
        total = sum(d['cost'] for d in daily_costs)
        avg = total / len(daily_costs) if daily_costs else 0
        max_d = max(daily_costs, key=lambda x: x['cost'])
        min_d = min(daily_costs, key=lambda x: x['cost'])
        over_budget_days = sum(1 for d in daily_costs if d['cost'] > daily_budget)
        
        stats = [
            f'Total Days Analyzed: {len(daily_costs)}',
            f'Total Spending: ${total:,.2f}',
            f'Average Daily Spending: ${avg:,.2f}',
            f'Maximum Daily Spending: ${max_d["cost"]:,.2f} ({max_d["date"]})',
            f'Minimum Daily Spending: ${min_d["cost"]:,.2f} ({min_d["date"]})',
            f'Days Over Budget: {over_budget_days}',
            f'Days Under Budget: {len(daily_costs) - over_budget_days}'
        ]
        
        for stat in stats:
            para = doc.add_paragraph()
            para.add_run('• ' + stat)


def analyze_daily_service_cost(svc):
    """Analyze service cost details to determine likely cause."""
    details = svc.get('details', [])
    if not details:
        return 'Insufficient detail data to determine specific cause.'
    
    # Sort by cost
    sorted_details = sorted(details, key=lambda x: x['cost'], reverse=True)
    top_usage = sorted_details[0] if sorted_details else None
    
    if not top_usage:
        return 'Unable to determine specific usage pattern.'
    
    usage_type = top_usage['usage_type']
    
    # Analyze based on service name and usage type
    service_name = svc['service'].lower()
    
    if 'ec2' in service_name:
        if 'boxusage' in usage_type.lower():
            return f'High compute usage detected ({simplify_usage_type(usage_type)}). Review instance sizes and running hours.'
        elif 'datatransfer' in usage_type.lower():
            return 'Data transfer costs are significant. Consider optimizing data movement patterns.'
        elif 'ebs' in usage_type.lower():
            return 'EBS storage costs detected. Review volume sizes and snapshot policies.'
        else:
            return f'Review EC2 usage pattern: {simplify_usage_type(usage_type)}'
    
    elif 's3' in service_name:
        if 'storage' in usage_type.lower():
            return 'S3 storage costs are high. Consider lifecycle policies and storage class optimization.'
        elif 'request' in usage_type.lower():
            return 'High number of S3 requests. Review access patterns and consider caching.'
        else:
            return f'Review S3 usage: {simplify_usage_type(usage_type)}'
    
    elif 'rds' in service_name:
        return 'Database costs detected. Review instance size, multi-AZ configuration, and backup retention.'
    
    elif 'lambda' in service_name:
        return 'Lambda costs high. Review function memory settings and invocation frequency.'
    
    elif 'cloudwatch' in service_name:
        return 'CloudWatch costs detected. Review metric collection, log retention, and alarm configurations.'
    
    else:
        return f'Primary cost driver: {simplify_usage_type(usage_type)}. Review service configuration and usage.'


def setup_document(doc):
    """Configure document settings, margins, and styles."""
    # Set page margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
    
    # Configure styles
    styles = doc.styles
    
    # Normal style
    normal = styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    
    # Heading 1
    h1 = styles['Heading 1']
    h1.font.name = 'Calibri Light'
    h1.font.size = Pt(24)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 51, 102)
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.keep_with_next = True
    
    # Heading 2
    h2 = styles['Heading 2']
    h2.font.name = 'Calibri Light'
    h2.font.size = Pt(16)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 82, 147)
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(8)
    h2.paragraph_format.keep_with_next = True
    
    # Heading 3
    h3 = styles['Heading 3']
    h3.font.name = 'Calibri'
    h3.font.size = Pt(13)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0, 102, 153)
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)
    h3.paragraph_format.keep_with_next = True


def add_cover_page(doc, month_names, budget_amount, breach_date):
    """Add a professional cover page."""
    # Add spacing at top
    for _ in range(4):
        doc.add_paragraph()
    
    # Company name
    company = doc.add_paragraph()
    company.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = company.add_run('ExamOnline')
    run.font.name = 'Calibri Light'
    run.font.size = Pt(42)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    # Decorative line
    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line.add_run('━' * 30)
    run.font.color.rgb = RGBColor(0, 102, 153)
    run.font.size = Pt(14)
    
    # Report title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('AWS Budget Breach')
    run.font.name = 'Calibri Light'
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(68, 68, 68)
    
    title2 = doc.add_paragraph()
    title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title2.add_run('Analysis Report')
    run.font.name = 'Calibri Light'
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(68, 68, 68)
    
    doc.add_paragraph()
    
    # Analysis period
    period = doc.add_paragraph()
    period.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = period.add_run('Analysis Period')
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(102, 102, 102)
    
    period_dates = doc.add_paragraph()
    period_dates.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = period_dates.add_run(f'{month_names[0]} — {month_names[-1]}')
    run.font.name = 'Calibri'
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    doc.add_paragraph()
    
    # Budget threshold (if provided)
    if budget_amount > 0:
        budget_label = doc.add_paragraph()
        budget_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = budget_label.add_run('Budget Threshold Exceeded')
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(153, 0, 0)
        
        budget_value = doc.add_paragraph()
        budget_value.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = budget_value.add_run(f'${budget_amount:,.2f}')
        run.font.name = 'Calibri'
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = RGBColor(153, 0, 0)
    
    # Add spacing before footer
    for _ in range(6):
        doc.add_paragraph()
    
    # Report date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(f'Report Generated: {datetime.now().strftime("%B %d, %Y")}')
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = RGBColor(102, 102, 102)
    
    # Confidential notice
    conf = doc.add_paragraph()
    conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    conf.paragraph_format.space_before = Pt(24)
    run = conf.add_run('CONFIDENTIAL')
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(153, 0, 0)
    
    doc.add_page_break()


def add_table_of_contents(doc):
    """Add a table of contents page."""
    toc_heading = doc.add_paragraph()
    run = toc_heading.add_run('Table of Contents')
    run.font.name = 'Calibri Light'
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    toc_heading.paragraph_format.space_after = Pt(24)
    
    # TOC entries
    toc_entries = [
        ('1.', 'Executive Summary', '3'),
        ('2.', 'Month-to-Date Cost Trends (Since Budget Reset)', '4'),
        ('3.', 'Cost Drivers Analysis', '5'),
        ('4.', 'Detailed Service Analysis', '6'),
        ('5.', 'Regional Cost Analysis', '8'),
        ('6.', 'Recommendations', '9'),
        ('7.', 'Appendix: Complete Data', '11'),
    ]
    
    for num, title, page in toc_entries:
        entry = doc.add_paragraph()
        entry.paragraph_format.space_after = Pt(12)
        
        # Number
        run = entry.add_run(f'{num}  ')
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)
        
        # Title
        run = entry.add_run(title)
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
        
        # Dots and page number
        run = entry.add_run('  ' + '.' * 60 + '  ')
        run.font.color.rgb = RGBColor(180, 180, 180)
        run.font.size = Pt(10)
        
        run = entry.add_run(page)
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
    
    doc.add_page_break()


def add_executive_summary(doc, increased_services, month_names, budget_amount,
                          overall_previous, overall_current, total_increase,
                          mtd_total=0, days_elapsed=0):
    """Add executive summary section."""
    doc.add_heading('Executive Summary', level=1)
    
    # Overview box
    add_info_box(doc, 'OVERVIEW', 
        f'This report provides a comprehensive analysis of AWS cost increases between '
        f'{month_names[0]} (baseline) and {month_names[-1]} (breach period). The analysis '
        f'compares month-over-month spending changes from when the budget resets on the 1st, '
        f'identifies services with cost growth, and provides actionable recommendations.',
        RGBColor(232, 245, 253))
    
    doc.add_paragraph()
    
    # Key Metrics Section
    doc.add_heading('Key Financial Metrics', level=2)
    
    # Row 1 - Labels
    labels = [f'{month_names[0]}', f'{month_names[-1]}', 'Month Change', 'MTD Spend', 'Services Impacted']
    
    # Create metrics table - columns based on labels count
    metrics_table = doc.add_table(rows=2, cols=len(labels))
    metrics_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    format_metrics_table(metrics_table)
    
    for i, label in enumerate(labels):
        cell = metrics_table.rows[0].cells[i]
        cell.text = label
        format_cell(cell, bold=True, bg_color='003366', font_color='FFFFFF', 
                   font_size=9, align='center')
    
    # Row 2 - Values
    overall_change = overall_current - overall_previous
    change_sign = '+' if overall_change >= 0 else ''
    mtd_display = f'${mtd_total:,.2f}' if mtd_total > 0 else f'${overall_current:,.2f}'
    values = [
        f'${overall_previous:,.2f}',
        f'${overall_current:,.2f}',
        f'{change_sign}${overall_change:,.2f}',
        mtd_display,
        str(len(increased_services))
    ]
    for i, value in enumerate(values):
        cell = metrics_table.rows[1].cells[i]
        cell.text = value
        if i == 2 and overall_change > 0:
            bg = 'FFE6E6'
        elif i == 3 and budget_amount > 0 and mtd_total > budget_amount:
            bg = 'FFE6E6'
        else:
            bg = 'F5F5F5'
        format_cell(cell, bold=True, bg_color=bg, font_size=12, align='center')
    
    doc.add_paragraph()
    
    # Budget Status (if provided)
    if budget_amount > 0:
        doc.add_heading('Budget Status', level=2)
        
        overage = overall_current - budget_amount
        overage_pct = (overage / budget_amount * 100) if budget_amount > 0 else 0
        
        # Calculate daily burn rate with actual days in month
        current_month_dt = datetime.strptime(month_names[-1], "%B %Y")
        days_in_month = calendar.monthrange(current_month_dt.year, current_month_dt.month)[1]
        
        if days_elapsed > 0:
            daily_avg = mtd_total / days_elapsed
            projected_month_end = daily_avg * days_in_month
        else:
            daily_avg = overall_current / days_in_month
            projected_month_end = overall_current
        
        if overage > 0:
            add_alert_box(doc, '⚠️ BUDGET EXCEEDED',
                f'Current spending of ${overall_current:,.2f} has exceeded the budget '
                f'threshold of ${budget_amount:,.2f} by ${overage:,.2f} ({overage_pct:.1f}%). '
                f'Average daily spend: ${daily_avg:,.2f}. '
                f'Immediate action is required to identify and address cost drivers.',
                RGBColor(255, 235, 235), RGBColor(153, 0, 0))
        else:
            add_alert_box(doc, '✓ WITHIN BUDGET',
                f'Current spending of ${overall_current:,.2f} is within the budget '
                f'threshold of ${budget_amount:,.2f}. '
                f'Projected month-end: ${projected_month_end:,.2f}.',
                RGBColor(235, 255, 235), RGBColor(0, 102, 0))
        
        doc.add_paragraph()
    
    # Top Cost Drivers Summary
    doc.add_heading('Top 5 Cost Increase Drivers', level=2)
    
    if increased_services[:5]:
        summary_table = doc.add_table(rows=6, cols=5)
        summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        format_data_table(summary_table)
        
        # Headers
        headers = ['Service', month_names[0], month_names[-1], 'Increase', '% Change']
        for i, header in enumerate(headers):
            cell = summary_table.rows[0].cells[i]
            cell.text = header
            format_cell(cell, bold=True, bg_color='0052CC', font_color='FFFFFF',
                       font_size=10, align='center')
        
        # Data rows
        for row_idx, svc in enumerate(increased_services[:5], 1):
            row = summary_table.rows[row_idx]
            
            row.cells[0].text = truncate_service_name(svc['service'])
            format_cell(row.cells[0], font_size=10, align='left')
            
            row.cells[1].text = f"${svc['previous_cost']:,.2f}"
            format_cell(row.cells[1], font_size=10, align='right')
            
            row.cells[2].text = f"${svc['current_cost']:,.2f}"
            format_cell(row.cells[2], font_size=10, align='right')
            
            row.cells[3].text = f"${svc['change']:,.2f}"
            format_cell(row.cells[3], font_size=10, align='right', bg_color='FFEEEE')
            
            row.cells[4].text = f"{svc['pct_change']:.1f}%"
            bg = 'FF6666' if svc['pct_change'] > 50 else 'FFCCCC' if svc['pct_change'] > 20 else 'FFE6E6'
            format_cell(row.cells[4], font_size=10, align='center', bg_color=bg, bold=True)
    
    doc.add_page_break()


def add_daily_cost_trends(doc, daily_costs, daily_service_costs, budget_amount, 
                          current_month_name, increased_services):
    """Add daily cost trends section showing MTD spending from budget reset."""
    doc.add_heading('Month-to-Date Cost Trends (Since Budget Reset)', level=1)
    
    intro = doc.add_paragraph()
    intro.add_run(
        f'This section shows daily cost accumulation for {current_month_name}, '
        f'starting from the 1st when the monthly budget resets. '
        f'Understanding daily spend patterns helps identify when cost spikes occurred '
        f'and which services contributed to the budget breach.'
    )
    intro.paragraph_format.space_after = Pt(16)
    
    # Calculate MTD statistics
    if daily_costs:
        total_mtd = sum(d['cost'] for d in daily_costs)
        days = len(daily_costs)
        avg_daily = total_mtd / days if days > 0 else 0
        max_day = max(daily_costs, key=lambda x: x['cost'])
        min_day = min(daily_costs, key=lambda x: x['cost'])
        
        # MTD Summary Box
        doc.add_heading('MTD Spending Summary', level=2)
        
        summary_table = doc.add_table(rows=2, cols=4)
        summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        format_metrics_table(summary_table)
        
        labels = ['Days Tracked', 'Total MTD Spend', 'Avg Daily Spend', 'Budget Used']
        for i, label in enumerate(labels):
            cell = summary_table.rows[0].cells[i]
            cell.text = label
            format_cell(cell, bold=True, bg_color='003366', font_color='FFFFFF',
                       font_size=9, align='center')
        
        budget_used_pct = (total_mtd / budget_amount * 100) if budget_amount > 0 else 0
        values = [
            str(days),
            f'${total_mtd:,.2f}',
            f'${avg_daily:,.2f}',
            f'{budget_used_pct:.1f}%' if budget_amount > 0 else 'N/A'
        ]
        for i, value in enumerate(values):
            cell = summary_table.rows[1].cells[i]
            cell.text = value
            bg = 'FFE6E6' if i == 3 and budget_used_pct > 100 else 'F5F5F5'
            format_cell(cell, bold=True, bg_color=bg, font_size=12, align='center')
        
        doc.add_paragraph()
        
        # Daily Breakdown Table
        doc.add_heading('Daily Cost Breakdown', level=2)
        
        # Show daily costs in a table (max 15 days to fit page)
        display_days = daily_costs[:MAX_DAILY_DISPLAY_DAYS] if len(daily_costs) > MAX_DAILY_DISPLAY_DAYS else daily_costs
        
        daily_table = doc.add_table(rows=len(display_days) + 1, cols=4)
        daily_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        format_data_table(daily_table)
        
        headers = ['Date', 'Daily Cost', 'Cumulative Total', 'vs Budget']
        for i, header in enumerate(headers):
            cell = daily_table.rows[0].cells[i]
            cell.text = header
            format_cell(cell, bold=True, bg_color='4A86C7', font_color='FFFFFF',
                       font_size=9, align='center')
        
        cumulative = 0
        for row_idx, day_data in enumerate(display_days, 1):
            row = daily_table.rows[row_idx]
            cumulative += day_data['cost']
            
            # Format date nicely
            date_obj = datetime.strptime(day_data['date'], '%Y-%m-%d')
            row.cells[0].text = date_obj.strftime('%b %d')
            format_cell(row.cells[0], font_size=9, align='center')
            
            row.cells[1].text = f"${day_data['cost']:,.2f}"
            format_cell(row.cells[1], font_size=9, align='right')
            
            row.cells[2].text = f"${cumulative:,.2f}"
            format_cell(row.cells[2], font_size=9, align='right')
            
            if budget_amount > 0:
                pct_of_budget = (cumulative / budget_amount * 100)
                row.cells[3].text = f"{pct_of_budget:.1f}%"
                bg = 'FF6666' if pct_of_budget > 100 else 'FFCCCC' if pct_of_budget > 80 else 'F5F5F5'
            else:
                row.cells[3].text = 'N/A'
                bg = 'F5F5F5'
            format_cell(row.cells[3], font_size=9, align='center', bg_color=bg)
        
        doc.add_paragraph()
        
        # Peak Spending Days
        doc.add_heading('Peak Spending Analysis', level=2)
        
        peak_para = doc.add_paragraph()
        peak_para.add_run('Highest Spend Day: ').bold = True
        max_date = datetime.strptime(max_day['date'], '%Y-%m-%d')
        peak_para.add_run(f"{max_date.strftime('%B %d')} - ${max_day['cost']:,.2f}")
        
        low_para = doc.add_paragraph()
        low_para.add_run('Lowest Spend Day: ').bold = True
        min_date = datetime.strptime(min_day['date'], '%Y-%m-%d')
        low_para.add_run(f"{min_date.strftime('%B %d')} - ${min_day['cost']:,.2f}")
        
        variance = max_day['cost'] - min_day['cost']
        variance_para = doc.add_paragraph()
        variance_para.add_run('Daily Variance: ').bold = True
        variance_para.add_run(f"${variance:,.2f} (indicates cost volatility)")
        
        doc.add_paragraph()
        
        # Top Service Contributors in Current Month
        if daily_service_costs:
            doc.add_heading('Top Daily Cost Contributors', level=2)
            
            # Calculate total for each service
            service_totals = []
            for service, costs in daily_service_costs.items():
                total = sum(c['cost'] for c in costs)
                if total > 0:
                    service_totals.append({'service': service, 'total': total})
            
            service_totals.sort(key=lambda x: x['total'], reverse=True)
            top_services = service_totals[:5]
            
            if top_services:
                svc_table = doc.add_table(rows=len(top_services) + 1, cols=3)
                svc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                format_data_table(svc_table)
                
                headers = ['Service', 'MTD Total', '% of MTD Spend']
                for i, header in enumerate(headers):
                    cell = svc_table.rows[0].cells[i]
                    cell.text = header
                    format_cell(cell, bold=True, bg_color='996600', font_color='FFFFFF',
                               font_size=9, align='center')
                
                for row_idx, svc in enumerate(top_services, 1):
                    row = svc_table.rows[row_idx]
                    
                    row.cells[0].text = truncate_service_name(svc['service'])
                    format_cell(row.cells[0], font_size=9, align='left')
                    
                    row.cells[1].text = f"${svc['total']:,.2f}"
                    format_cell(row.cells[1], font_size=9, align='right')
                    
                    pct = (svc['total'] / total_mtd * 100) if total_mtd > 0 else 0
                    row.cells[2].text = f"{pct:.1f}%"
                    format_cell(row.cells[2], font_size=9, align='center')
    
    doc.add_page_break()


def add_cost_drivers_analysis(doc, increased_services, month_names, total_increase):
    """Add detailed cost drivers analysis."""
    doc.add_heading('Cost Drivers Analysis', level=1)
    
    intro = doc.add_paragraph()
    intro.add_run(
        'This section provides an in-depth analysis of the primary factors driving '
        'AWS cost increases. Each service is examined to identify root causes and '
        'quantify its contribution to the overall cost growth.'
    )
    intro.paragraph_format.space_after = Pt(16)
    
    # Contribution Analysis
    doc.add_heading('Cost Increase Contribution by Service', level=2)
    
    if increased_services and total_increase > 0:
        # Show top 8 contributors
        top_contributors = increased_services[:8]
        
        contrib_table = doc.add_table(rows=len(top_contributors) + 1, cols=4)
        contrib_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        format_data_table(contrib_table)
        
        headers = ['Service', 'Cost Increase', '% of Total Increase', 'Impact Level']
        for i, header in enumerate(headers):
            cell = contrib_table.rows[0].cells[i]
            cell.text = header
            format_cell(cell, bold=True, bg_color='003366', font_color='FFFFFF',
                       font_size=10, align='center')
        
        for row_idx, svc in enumerate(top_contributors, 1):
            row = contrib_table.rows[row_idx]
            contribution = (svc['change'] / total_increase * 100) if total_increase > 0 else 0
            
            # Determine impact level
            if contribution > 30:
                impact = 'CRITICAL'
                impact_color = 'FF0000'
            elif contribution > 15:
                impact = 'HIGH'
                impact_color = 'FF6600'
            elif contribution > 5:
                impact = 'MEDIUM'
                impact_color = 'FFCC00'
            else:
                impact = 'LOW'
                impact_color = '00CC00'
            
            row.cells[0].text = truncate_service_name(svc['service'])
            format_cell(row.cells[0], font_size=10, align='left')
            
            row.cells[1].text = f"${svc['change']:,.2f}"
            format_cell(row.cells[1], font_size=10, align='right')
            
            row.cells[2].text = f"{contribution:.1f}%"
            format_cell(row.cells[2], font_size=10, align='center')
            
            row.cells[3].text = impact
            format_cell(row.cells[3], font_size=9, align='center', bold=True,
                       font_color=impact_color)
    
    doc.add_paragraph()
    
    # Primary Drivers Detail
    doc.add_heading('Primary Cost Drivers - Detailed Analysis', level=2)
    
    for i, svc in enumerate(increased_services[:5], 1):
        contribution = (svc['change'] / total_increase * 100) if total_increase > 0 else 0
        
        # Service heading
        svc_para = doc.add_paragraph()
        svc_para.paragraph_format.space_before = Pt(16)
        run = svc_para.add_run(f'{i}. {svc["service"]}')
        run.font.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0, 51, 102)
        
        # Quick stats in a mini table
        stats_table = doc.add_table(rows=1, cols=4)
        format_inline_stats_table(stats_table)
        
        stats = [
            f"Previous: ${svc['previous_cost']:,.2f}",
            f"Current: ${svc['current_cost']:,.2f}",
            f"Increase: ${svc['change']:,.2f}",
            f"Growth: {svc['pct_change']:.1f}%"
        ]
        for j, stat in enumerate(stats):
            stats_table.rows[0].cells[j].text = stat
            format_cell(stats_table.rows[0].cells[j], font_size=9, align='center',
                       bg_color='F0F5FF')
        
        # Analysis
        driver_analysis = analyze_service_drivers(svc)
        
        analysis_para = doc.add_paragraph()
        analysis_para.paragraph_format.left_indent = Inches(0.25)
        
        run = analysis_para.add_run('Root Cause: ')
        run.font.bold = True
        run.font.size = Pt(10)
        analysis_para.add_run(driver_analysis['primary_driver'])
        
        if driver_analysis['usage_changes']:
            changes_para = doc.add_paragraph()
            changes_para.paragraph_format.left_indent = Inches(0.25)
            run = changes_para.add_run('Key Observations:')
            run.font.bold = True
            run.font.size = Pt(10)
            
            for change in driver_analysis['usage_changes'][:3]:
                bullet = doc.add_paragraph(f'• {change}')
                bullet.paragraph_format.left_indent = Inches(0.5)
                bullet.paragraph_format.space_after = Pt(2)
    
    doc.add_page_break()


def add_detailed_service_analysis(doc, increased_services, month_names, months):
    """Add detailed breakdown for each service."""
    doc.add_heading('Detailed Service Analysis', level=1)
    
    intro = doc.add_paragraph()
    intro.add_run(
        'This section provides a granular breakdown of cost increases for each '
        'impacted AWS service, including specific usage types and resource categories.'
    )
    intro.paragraph_format.space_after = Pt(16)
    
    for svc in increased_services[:10]:  # Top 10 services
        doc.add_heading(truncate_service_name(svc['service']), level=2)
        
        # Cost summary box
        summary = doc.add_paragraph()
        run = summary.add_run(f"${svc['previous_cost']:,.2f}")
        run.font.bold = True
        summary.add_run(f" ({month_names[0]})  →  ")
        run = summary.add_run(f"${svc['current_cost']:,.2f}")
        run.font.bold = True
        summary.add_run(f" ({month_names[-1]})  |  ")
        run = summary.add_run(f"Increase: ${svc['change']:,.2f} ({svc['pct_change']:.1f}%)")
        run.font.bold = True
        run.font.color.rgb = RGBColor(153, 0, 0)
        
        # Usage breakdown
        data = svc['data']
        current_details = data.get(months[-1], {}).get('details', [])
        previous_details = data.get(months[0], {}).get('details', [])
        prev_lookup = {d['usage_type']: d for d in previous_details}
        
        # Find increased usage types
        changes = []
        for detail in current_details:
            usage_type = detail['usage_type']
            current_cost = detail['cost']
            previous_cost = prev_lookup.get(usage_type, {}).get('cost', 0)
            change = current_cost - previous_cost
            if change > 0.01:
                changes.append({
                    'usage_type': usage_type,
                    'previous': previous_cost,
                    'current': current_cost,
                    'change': change
                })
        
        changes.sort(key=lambda x: x['change'], reverse=True)
        
        if changes[:5]:
            doc.add_heading('Usage Type Breakdown', level=3)
            
            usage_table = doc.add_table(rows=min(6, len(changes[:5]) + 1), cols=4)
            format_data_table(usage_table)
            
            headers = ['Usage Type', month_names[0], month_names[-1], 'Change']
            for i, header in enumerate(headers):
                cell = usage_table.rows[0].cells[i]
                cell.text = header
                format_cell(cell, bold=True, bg_color='4A86C7', font_color='FFFFFF',
                           font_size=9, align='center')
            
            for row_idx, change in enumerate(changes[:5], 1):
                row = usage_table.rows[row_idx]
                row.cells[0].text = simplify_usage_type(change['usage_type'])
                format_cell(row.cells[0], font_size=9, align='left')
                
                row.cells[1].text = f"${change['previous']:,.2f}"
                format_cell(row.cells[1], font_size=9, align='right')
                
                row.cells[2].text = f"${change['current']:,.2f}"
                format_cell(row.cells[2], font_size=9, align='right')
                
                row.cells[3].text = f"+${change['change']:,.2f}"
                format_cell(row.cells[3], font_size=9, align='right', bg_color='FFEEEE')
        
        # Reason analysis
        doc.add_heading('Analysis & Root Cause', level=3)
        reason = generate_detailed_reason(svc, changes)
        reason_para = doc.add_paragraph(reason)
        reason_para.paragraph_format.left_indent = Inches(0.25)
        
        doc.add_paragraph()


def add_regional_analysis(doc, regional_costs, months, month_names):
    """Add regional cost breakdown."""
    doc.add_heading('Regional Cost Analysis', level=1)
    
    intro = doc.add_paragraph()
    intro.add_run(
        'This section analyzes cost distribution across AWS regions to identify '
        'geographic patterns in cost increases.'
    )
    intro.paragraph_format.space_after = Pt(16)
    
    # Find regions with increases
    region_changes = []
    for region, costs in regional_costs.items():
        prev = costs.get(months[0], 0)
        curr = costs.get(months[-1], 0)
        change = curr - prev
        if change > 0:
            region_changes.append({
                'region': region,
                'previous': prev,
                'current': curr,
                'change': change
            })
    
    region_changes.sort(key=lambda x: x['change'], reverse=True)
    
    if region_changes:
        doc.add_heading('Regions with Cost Increases', level=2)
        
        table = doc.add_table(rows=len(region_changes) + 1, cols=4)
        format_data_table(table)
        
        headers = ['Region', month_names[0], month_names[-1], 'Increase']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            format_cell(cell, bold=True, bg_color='996600', font_color='FFFFFF',
                       font_size=10, align='center')
        
        for row_idx, rc in enumerate(region_changes, 1):
            row = table.rows[row_idx]
            row.cells[0].text = rc['region']
            format_cell(row.cells[0], font_size=10, align='left')
            
            row.cells[1].text = f"${rc['previous']:,.2f}"
            format_cell(row.cells[1], font_size=10, align='right')
            
            row.cells[2].text = f"${rc['current']:,.2f}"
            format_cell(row.cells[2], font_size=10, align='right')
            
            row.cells[3].text = f"+${rc['change']:,.2f}"
            format_cell(row.cells[3], font_size=10, align='right', bg_color='FFF5E6')
    else:
        doc.add_paragraph('No regions with significant cost increases were identified.')
    
    doc.add_page_break()


def add_recommendations(doc, increased_services):
    """Add recommendations section."""
    doc.add_heading('Recommendations', level=1)
    
    intro = doc.add_paragraph()
    intro.add_run(
        'Based on the analysis of cost increases, the following recommendations '
        'are provided to optimize AWS spending and prevent future budget breaches.'
    )
    intro.paragraph_format.space_after = Pt(16)
    
    # Immediate Actions
    doc.add_heading('Immediate Actions (This Week)', level=2)
    
    immediate = generate_immediate_actions(increased_services)
    for i, action in enumerate(immediate, 1):
        action_para = doc.add_paragraph()
        run = action_para.add_run(f'{i}. {action["title"]}')
        run.font.bold = True
        run.font.color.rgb = RGBColor(153, 0, 0)
        
        desc_para = doc.add_paragraph(action['description'])
        desc_para.paragraph_format.left_indent = Inches(0.25)
        
        if action.get('steps'):
            for step in action['steps']:
                step_para = doc.add_paragraph(f'• {step}')
                step_para.paragraph_format.left_indent = Inches(0.5)
                step_para.paragraph_format.space_after = Pt(2)
    
    doc.add_paragraph()
    
    # Short-term
    doc.add_heading('Short-term Optimizations (1-2 Weeks)', level=2)
    
    short_term = [
        ('Configure Budget Alerts', 
         'Set up AWS Budgets with alerts at 50%, 75%, and 90% thresholds for early warning.'),
        ('Enable Cost Allocation Tags', 
         'Implement mandatory tagging for all resources to enable detailed cost attribution.'),
        ('Review Trusted Advisor', 
         'Check AWS Trusted Advisor for immediate cost optimization recommendations.'),
        ('Audit Unused Resources', 
         'Identify and terminate unused EC2 instances, EBS volumes, and other idle resources.'),
    ]
    
    for title, desc in short_term:
        para = doc.add_paragraph()
        run = para.add_run(f'• {title}: ')
        run.font.bold = True
        para.add_run(desc)
    
    doc.add_paragraph()
    
    # Long-term
    doc.add_heading('Long-term Strategy', level=2)
    
    long_term = [
        ('Implement Savings Plans', 
         'Evaluate workload patterns and purchase Savings Plans for predictable compute usage '
         '(up to 72% savings).'),
        ('Enable Cost Anomaly Detection', 
         'Set up AWS Cost Anomaly Detection for automated alerts on unusual spending patterns.'),
        ('Establish Governance', 
         'Create policies for resource provisioning and establish regular cost review meetings.'),
        ('Right-size Resources', 
         'Use AWS Compute Optimizer to identify and right-size over-provisioned resources.'),
    ]
    
    for title, desc in long_term:
        para = doc.add_paragraph()
        run = para.add_run(f'• {title}: ')
        run.font.bold = True
        para.add_run(desc)
    
    doc.add_page_break()


def add_appendix(doc, increased_services, month_names):
    """Add appendix with complete data."""
    doc.add_heading('Appendix: Complete Cost Increase Data', level=1)
    
    doc.add_paragraph(
        'This appendix contains the complete list of all AWS services with cost '
        'increases during the analysis period, sorted by increase amount.'
    )
    
    if increased_services:
        table = doc.add_table(rows=len(increased_services) + 2, cols=5)
        format_data_table(table)
        
        # Headers
        headers = ['Service', month_names[0], month_names[-1], 'Increase', '% Change']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            format_cell(cell, bold=True, bg_color='333333', font_color='FFFFFF',
                       font_size=9, align='center')
        
        # Data
        total_prev = 0
        total_curr = 0
        total_change = 0
        
        for row_idx, svc in enumerate(increased_services, 1):
            row = table.rows[row_idx]
            
            row.cells[0].text = truncate_service_name(svc['service'])
            format_cell(row.cells[0], font_size=9, align='left')
            
            row.cells[1].text = f"${svc['previous_cost']:,.2f}"
            format_cell(row.cells[1], font_size=9, align='right')
            
            row.cells[2].text = f"${svc['current_cost']:,.2f}"
            format_cell(row.cells[2], font_size=9, align='right')
            
            row.cells[3].text = f"${svc['change']:,.2f}"
            format_cell(row.cells[3], font_size=9, align='right')
            
            row.cells[4].text = f"{svc['pct_change']:.1f}%"
            format_cell(row.cells[4], font_size=9, align='center')
            
            total_prev += svc['previous_cost']
            total_curr += svc['current_cost']
            total_change += svc['change']
        
        # Totals row
        totals_row = table.rows[-1]
        totals_row.cells[0].text = 'TOTAL'
        format_cell(totals_row.cells[0], bold=True, bg_color='333333', 
                   font_color='FFFFFF', font_size=9, align='left')
        
        totals_row.cells[1].text = f"${total_prev:,.2f}"
        format_cell(totals_row.cells[1], bold=True, bg_color='333333',
                   font_color='FFFFFF', font_size=9, align='right')
        
        totals_row.cells[2].text = f"${total_curr:,.2f}"
        format_cell(totals_row.cells[2], bold=True, bg_color='333333',
                   font_color='FFFFFF', font_size=9, align='right')
        
        totals_row.cells[3].text = f"${total_change:,.2f}"
        format_cell(totals_row.cells[3], bold=True, bg_color='333333',
                   font_color='FFFFFF', font_size=9, align='right')
        
        totals_row.cells[4].text = ''
        format_cell(totals_row.cells[4], bg_color='333333')


# ===== HELPER FUNCTIONS =====

def format_cell(cell, bold=False, bg_color=None, font_color=None, 
                font_size=11, align='left'):
    """Format a table cell with consistent styling."""
    # Set vertical alignment
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
    # Format paragraph
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(4)
        paragraph.paragraph_format.space_after = Pt(4)
        
        if align == 'center':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == 'right':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        for run in paragraph.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(font_size)
            run.font.bold = bold
            if font_color:
                # Parse hex color string to RGB
                r = int(font_color[0:2], 16)
                g = int(font_color[2:4], 16)
                b = int(font_color[4:6], 16)
                run.font.color.rgb = RGBColor(r, g, b)
    
    # Set background color
    if bg_color:
        set_cell_shading(cell, bg_color)


def set_cell_shading(cell, color):
    """Set background color for a table cell."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def format_data_table(table):
    """Apply consistent formatting to a data table."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Set table borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), 'CCCCCC')
        tblBorders.append(border)
    
    tblPr.append(tblBorders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)


def format_metrics_table(table):
    """Format the key metrics table."""
    format_data_table(table)
    
    # Set column widths
    for row in table.rows:
        for cell in row.cells:
            cell.width = Inches(1.5)


def format_inline_stats_table(table):
    """Format inline statistics table."""
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    
    # Remove borders for inline stats
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '2')
        border.set(qn('w:color'), 'DDDDDD')
        tblBorders.append(border)
    
    tblPr.append(tblBorders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)


def add_info_box(doc, title, content, bg_color):
    """Add an information box with background."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    
    # Title
    title_para = cell.paragraphs[0]
    run = title_para.add_run(title)
    run.font.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    # Content
    content_para = cell.add_paragraph(content)
    content_para.paragraph_format.space_before = Pt(8)
    
    # Background
    set_cell_shading(cell, 'E8F4FC')
    
    # Padding
    cell.paragraphs[0].paragraph_format.space_before = Pt(12)
    cell.paragraphs[-1].paragraph_format.space_after = Pt(12)
    
    for para in cell.paragraphs:
        para.paragraph_format.left_indent = Inches(0.15)
        para.paragraph_format.right_indent = Inches(0.15)


def add_alert_box(doc, title, content, bg_color, title_color):
    """Add an alert box.
    
    Args:
        doc: Document object
        title: Box title text
        content: Box content text
        bg_color: Background color (RGBColor object)
        title_color: Title text color (RGBColor object)
    """
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    
    # Title
    title_para = cell.paragraphs[0]
    run = title_para.add_run(title)
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = title_color
    
    # Content
    content_para = cell.add_paragraph(content)
    content_para.paragraph_format.space_before = Pt(8)
    
    # Background color - RGBColor object converts to hex string
    hex_color = str(bg_color)
    set_cell_shading(cell, hex_color)
    
    # Padding
    cell.paragraphs[0].paragraph_format.space_before = Pt(12)
    cell.paragraphs[-1].paragraph_format.space_after = Pt(12)
    
    for para in cell.paragraphs:
        para.paragraph_format.left_indent = Inches(0.15)
        para.paragraph_format.right_indent = Inches(0.15)


def truncate_service_name(name, max_len=40):
    """Truncate long service names."""
    if len(name) <= max_len:
        return name
    return name[:max_len-3] + '...'


def simplify_usage_type(usage_type):
    """Simplify AWS usage type names."""
    if ':' in usage_type:
        parts = usage_type.split(':')
        if len(parts) > 1 and parts[-1]:
            result = parts[-1]
            if len(result) > 35:
                return result[:32] + '...'
            return result
    if len(usage_type) > 35:
        return usage_type[:32] + '...'
    return usage_type


def analyze_service_drivers(svc):
    """Analyze cost drivers for a service."""
    service = svc['service']
    
    result = {
        'primary_driver': 'Increased service usage and resource consumption',
        'usage_changes': [],
        'recommendations': []
    }
    
    if 'EC2' in service:
        result['primary_driver'] = 'Increased compute instance usage, instance type scaling, or extended runtime hours'
        result['usage_changes'] = [
            'Additional EC2 instances provisioned or upgraded to larger instance types',
            'Instances running for extended hours or 24/7 instead of scheduled hours',
            'Data transfer costs associated with EC2 instances increased'
        ]
    elif 'S3' in service:
        result['primary_driver'] = 'Increased storage volume, data transfer, or API request volume'
        result['usage_changes'] = [
            'Storage volume growth from new data uploads',
            'Increased data transfer (cross-region or internet egress)',
            'Higher API request volume (GET, PUT, LIST operations)'
        ]
    elif 'RDS' in service:
        result['primary_driver'] = 'Database instance scaling, storage growth, or increased IOPS'
        result['usage_changes'] = [
            'Database instance upgraded to larger instance class',
            'Storage auto-scaling triggered by data growth',
            'Increased provisioned IOPS or Multi-AZ deployment'
        ]
    elif 'Lambda' in service:
        result['primary_driver'] = 'Increased function invocations, duration, or memory allocation'
        result['usage_changes'] = [
            'Higher number of function invocations from increased traffic',
            'Longer function execution times due to processing complexity',
            'Memory allocation increases affecting cost per invocation'
        ]
    elif 'CloudWatch' in service:
        result['primary_driver'] = 'Increased logging, metrics collection, or dashboard usage'
        result['usage_changes'] = [
            'Log ingestion volume increased significantly',
            'Additional custom metrics or high-resolution metrics enabled',
            'More frequent API calls or additional CloudWatch dashboards'
        ]
    elif 'Transfer' in service or 'CloudFront' in service:
        result['primary_driver'] = 'Increased data transfer volume or CDN usage'
        result['usage_changes'] = [
            'Higher data egress to internet or cross-region',
            'Increased CDN traffic and cache invalidations',
            'Cross-AZ or cross-region data movement increased'
        ]
    else:
        result['usage_changes'] = [
            'Service usage volume increased during the analysis period',
            'Resource configuration changes may have impacted costs',
            'API call volume or data processing increased'
        ]
    
    return result


def generate_detailed_reason(svc, changes):
    """Generate detailed reason for cost increase."""
    service = svc['service']
    pct = svc['pct_change']
    
    if pct > 100:
        severity = 'dramatic'
    elif pct > 50:
        severity = 'significant'
    elif pct > 20:
        severity = 'notable'
    else:
        severity = 'moderate'
    
    reason = f"The {severity} cost increase of {pct:.1f}% for {truncate_service_name(service)} "
    
    if changes:
        top = changes[0]
        reason += f"is primarily attributed to '{simplify_usage_type(top['usage_type'])}', "
        reason += f"which grew from ${top['previous']:,.2f} to ${top['current']:,.2f} "
        reason += f"(+${top['change']:,.2f}). "
    
    if 'EC2' in service:
        reason += "This typically indicates additional compute capacity, larger instances, or extended running hours. "
        reason += "Consider reviewing instance utilization and implementing auto-scaling or scheduled shutdowns."
    elif 'S3' in service:
        reason += "This suggests increased storage consumption or data transfer activity. "
        reason += "Review bucket sizes, implement lifecycle policies, and optimize data transfer patterns."
    elif 'RDS' in service:
        reason += "This may reflect database scaling, storage growth, or IOPS increases. "
        reason += "Evaluate instance right-sizing and consider Reserved Instances for stable workloads."
    elif 'Lambda' in service:
        reason += "This indicates higher function invocations or longer execution times. "
        reason += "Optimize function code and review memory allocation settings."
    else:
        reason += "Review the specific usage patterns for this service to identify optimization opportunities."
    
    return reason


def generate_immediate_actions(increased_services):
    """Generate immediate action recommendations."""
    actions = []
    
    ec2_services = [s for s in increased_services if 'EC2' in s['service']]
    if ec2_services:
        total = sum(s['change'] for s in ec2_services)
        actions.append({
            'title': 'Audit EC2 Instance Usage',
            'description': f'EC2-related costs increased by ${total:,.2f}. Perform immediate review.',
            'steps': [
                'Identify and terminate unused or idle instances',
                'Review instances running outside business hours',
                'Check for over-provisioned instance types'
            ]
        })
    
    storage_services = [s for s in increased_services if 'S3' in s['service'] or 'EBS' in s['service']]
    if storage_services:
        total = sum(s['change'] for s in storage_services)
        actions.append({
            'title': 'Review Storage Resources',
            'description': f'Storage costs increased by ${total:,.2f}. Audit storage utilization.',
            'steps': [
                'Delete unused EBS volumes and old snapshots',
                'Implement S3 lifecycle policies for data archival',
                'Review and remove unnecessary data'
            ]
        })
    
    if not actions and increased_services:
        top = increased_services[0]
        actions.append({
            'title': f'Investigate {truncate_service_name(top["service"])}',
            'description': f'Largest cost increase of ${top["change"]:,.2f}. Requires immediate review.',
            'steps': [
                'Review recent configuration changes',
                'Analyze usage patterns and trends',
                'Identify optimization opportunities'
            ]
        })
    
    return actions


# Compute usage type patterns
COMPUTE_PATTERNS = [
    ('BoxUsage:', 'EC2'),
    ('HeavyUsage:', 'EC2 Reserved'),
    ('SpotUsage:', 'EC2 Spot'),
    ('InstanceUsage:', 'RDS'),
    ('Multi-AZUsage:', 'RDS Multi-AZ'),
    ('ServerlessUsage:', 'RDS Serverless'),
    ('NodeUsage:', 'ElastiCache'),
    ('Node:', 'Redshift'),
]


def is_compute_usage_type(usage_type):
    """Check if usage type is compute resource."""
    for pattern, _ in COMPUTE_PATTERNS:
        if pattern in usage_type:
            return True
    return False
